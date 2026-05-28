"""Export documents, summaries, articles, or arbitrary chunk bundles
to Markdown / DOCX / PDF.

Three target formats:

- ``md`` — Markdown string / file. Headings, bullets, blockquotes.
- ``docx`` — Microsoft Word via python-docx. Headings → Word heading
  styles, code blocks → fixed-width.
- ``pdf`` — Lightweight ReportLab Platypus rendering. Headings,
  paragraphs, code, bullets. Not pixel-perfect — good enough for a
  shareable deliverable.

All exporters share a single intermediate representation (the
Markdown source). MD is the source of truth; DOCX/PDF render from it.
"""

from __future__ import annotations

import re
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from kglite_docs.store import Store
from kglite_docs.store import rows as _rows

# ─── source assembly ──────────────────────────────────────────────────────


def document_to_md(store: Store, doc_id: str, *, include_summaries: bool = False) -> str:
    """Render a stored document back to markdown, in reading order.
    Optionally inline `verified` summaries beneath each chunk."""
    head = _rows(store.cypher(
        "MATCH (d:Document {id: $id}) RETURN d.title AS title, d.path AS path",
        params={"id": doc_id},
    ))
    if not head:
        return ""
    parts: list[str] = [f"# {head[0]['title']}\n"]
    chunks = _rows(store.cypher(
        "MATCH (d:Document {id: $id})-[:HAS_CHUNK]->(c:Chunk) "
        "RETURN c.id AS id, c.page_number AS page, c.text AS text, "
        "c.headings_json AS headings ORDER BY c.page_number, c.chunk_index",
        params={"id": doc_id},
    ))
    last_page: int | None = None
    for c in chunks:
        page = c.get("page")
        if page != last_page:
            parts.append(f"\n---\n*Page {page}*\n")
            last_page = page
        parts.append(c.get("text") or "")
        if include_summaries:
            # Verified summaries via label predicate (kglite 0.10.5)
            sums = _rows(store.cypher(
                "MATCH (s:Summary:Verified) WHERE s.target_id = $cid "
                "RETURN s.text AS text",
                params={"cid": c["id"]},
            ))
            for s in sums:
                parts.append(f"\n> **Summary**: {s['text']}\n")
    return "\n\n".join(p for p in parts if p)


def summary_to_md(store: Store, summary_id: str) -> str:
    rows = _rows(store.cypher(
        "MATCH (s:Summary {id: $sid}) "
        "OPTIONAL MATCH (a:Agent)-[:AUTHORED]->(s) "
        "OPTIONAL MATCH (s)-[:VERIFIED_BY]->(v:Agent) "
        "RETURN s.text AS text, s.depth AS depth, s.verification_status AS status, "
        "s.created_at AS created_at, s.verified_at AS verified_at, "
        "a.id AS author, v.id AS verifier",
        params={"sid": summary_id},
    ))
    if not rows:
        return ""
    s = rows[0]
    md = [
        f"# Summary ({s.get('depth', 'chunk')})\n",
        s.get("text", ""),
        "",
        f"- **Author**: {s.get('author', 'unknown')}",
        f"- **Status**: {s.get('status', 'unverified')}",
    ]
    if s.get("verifier"):
        md.append(f"- **Verified by**: {s['verifier']} at {s.get('verified_at', '')}")
    return "\n".join(md)


def cluster_to_md(
    store: Store,
    cluster_id: str,
    *,
    include_member_text: bool = False,
) -> str:
    cl = _rows(store.cypher(
        "MATCH (cl:Cluster {id: $id}) "
        "RETURN cl.title AS title, cl.algorithm AS algorithm, cl.size AS size",
        params={"id": cluster_id},
    ))
    if not cl:
        return ""
    parts = [f"# Cluster: {cl[0].get('title', cluster_id)}\n"]
    parts.append(f"- **Algorithm**: {cl[0].get('algorithm')}")
    parts.append(f"- **Members**: {cl[0].get('size')}\n")
    members = _rows(store.cypher(
        "MATCH (c:Chunk)-[:IN_CLUSTER]->(cl:Cluster {id: $id}) "
        "OPTIONAL MATCH (s:Summary) WHERE s.target_id = c.id "
        "RETURN c.id AS id, c.doc_id AS doc_id, c.page_number AS page, "
        "c.text AS text, collect(s.text)[0] AS summary "
        "ORDER BY c.doc_id, c.page_number",
        params={"id": cluster_id},
    ))
    parts.append("## Members\n")
    for m in members:
        parts.append(f"### {m['id']}")
        if m.get("summary"):
            parts.append(f"> {m['summary']}")
        if include_member_text:
            parts.append((m.get("text") or "")[:1500])
        parts.append("")
    return "\n".join(parts)


# ─── format conversion ────────────────────────────────────────────────────


def md_to_docx(md: str, out_path: str | Path) -> Path:
    """Render markdown to .docx using python-docx."""
    from docx import Document  # type: ignore

    doc = Document()
    for line in md.splitlines():
        if not line.strip():
            doc.add_paragraph("")
            continue
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=min(level, 9))
            continue
        if line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.style = "Intense Quote"
            continue
        if line.startswith(("- ", "* ", "+ ")):
            doc.add_paragraph(line[2:], style="List Bullet")
            continue
        doc.add_paragraph(line)
    out = Path(out_path)
    doc.save(str(out))
    return out


def md_to_pdf(md: str, out_path: str | Path, *, title: str = "") -> Path:
    """Render markdown to .pdf using ReportLab Platypus. Lightweight,
    handles headings + paragraphs + bullets + blockquotes + code."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Paragraph,
        Preformatted,
        SimpleDocTemplate,
        Spacer,
    )

    styles = getSampleStyleSheet()
    h_styles = {
        i: ParagraphStyle(
            name=f"H{i}", parent=styles["Heading1" if i == 1 else "Heading2"],
            fontSize=max(20 - i * 2, 11), spaceAfter=8, spaceBefore=12,
        )
        for i in range(1, 7)
    }
    body = ParagraphStyle(
        name="Body", parent=styles["BodyText"], fontSize=10.5, leading=14, spaceAfter=6,
    )
    bullet = ParagraphStyle(name="Bullet", parent=body, leftIndent=18, bulletIndent=6)
    quote = ParagraphStyle(
        name="Quote", parent=body, leftIndent=18,
        textColor="#444", fontName="Helvetica-Oblique",
    )

    flow: list[Any] = []
    if title:
        flow.append(Paragraph(title, h_styles[1]))
        flow.append(Spacer(1, 0.1 * inch))

    in_code = False
    code_buf: list[str] = []
    for line in md.splitlines():
        if line.startswith("```"):
            if in_code:
                flow.append(Preformatted("\n".join(code_buf), styles["Code"]))
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        if not line.strip():
            flow.append(Spacer(1, 0.08 * inch))
            continue
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            flow.append(Paragraph(_escape(m.group(2)), h_styles[level]))
            continue
        if line.startswith("> "):
            flow.append(Paragraph(_escape(line[2:]), quote))
            continue
        if line.startswith(("- ", "* ", "+ ")):
            flow.append(Paragraph("• " + _escape(line[2:]), bullet))
            continue
        flow.append(Paragraph(_escape(line), body))

    out = Path(out_path)
    SimpleDocTemplate(
        str(out), pagesize=letter,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
    ).build(flow)
    return out


def _escape(text: str) -> str:
    """Escape minimal HTML/ReportLab-special characters in inline text.
    Markdown emphasis is left as plain characters for v1."""
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )


# ─── public entry points ──────────────────────────────────────────────────


def export_document(
    corpus: Any,
    doc_id: str,
    out_path: str | Path,
    *,
    format: str | None = None,
    include_summaries: bool = False,
) -> Path:
    """Export one document to MD / DOCX / PDF. Format inferred from
    `out_path` extension; pass `format=` to override."""
    return _export(
        md=document_to_md(corpus.store, doc_id, include_summaries=include_summaries),
        out_path=out_path, format=format, title=_doc_title(corpus, doc_id),
    )


def export_cluster(
    corpus: Any,
    cluster_id: str,
    out_path: str | Path,
    *,
    format: str | None = None,
    include_member_text: bool = False,
) -> Path:
    return _export(
        md=cluster_to_md(corpus.store, cluster_id, include_member_text=include_member_text),
        out_path=out_path, format=format, title=f"Cluster {cluster_id}",
    )


def export_summary(
    corpus: Any,
    summary_id: str,
    out_path: str | Path,
    *,
    format: str | None = None,
) -> Path:
    return _export(
        md=summary_to_md(corpus.store, summary_id),
        out_path=out_path, format=format, title="Summary",
    )


def export_bundle(
    corpus: Any,
    items: Iterable[dict[str, Any]],
    out_path: str | Path,
    *,
    format: str | None = None,
    title: str = "Synthesis bundle",
) -> Path:
    """Export an arbitrary set of items to a single MD/DOCX/PDF.

    Each `item` is a dict with one of:

    - ``{"kind": "doc", "id": ...}``
    - ``{"kind": "cluster", "id": ...}``
    - ``{"kind": "summary", "id": ...}``
    - ``{"kind": "markdown", "text": ...}`` — raw passthrough
    """
    md_parts: list[str] = [f"# {title}\n"]
    for it in items:
        kind = it.get("kind")
        if kind == "doc":
            md_parts.append(document_to_md(corpus.store, it["id"], include_summaries=True))
        elif kind == "cluster":
            md_parts.append(cluster_to_md(corpus.store, it["id"], include_member_text=False))
        elif kind == "summary":
            md_parts.append(summary_to_md(corpus.store, it["id"]))
        elif kind == "markdown":
            md_parts.append(it.get("text", ""))
        md_parts.append("\n---\n")
    return _export("\n\n".join(md_parts), out_path, format=format, title=title)


def _export(md: str, out_path: str | Path, *, format: str | None, title: str) -> Path:
    out = Path(out_path)
    fmt = (format or out.suffix.lstrip(".")).lower()
    if fmt in {"md", "markdown", ""}:
        out.write_text(md, encoding="utf-8")
        return out
    if fmt == "docx":
        return md_to_docx(md, out)
    if fmt == "pdf":
        return md_to_pdf(md, out, title=title)
    raise ValueError(f"unsupported export format: {fmt!r}")


def _doc_title(corpus: Any, doc_id: str) -> str:
    rows = _rows(corpus.cypher(
        "MATCH (d:Document {id: $id}) RETURN d.title AS title",
        params={"id": doc_id},
    ))
    return rows[0]["title"] if rows else doc_id
